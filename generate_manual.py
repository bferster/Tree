import docx
from docx.shared import Pt, Inches

doc = docx.Document()
doc.add_heading('Verité: Comprehensive User Manual', 0)

# ----------------- INTRODUCTION -----------------
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    "Welcome to the Verité application. Verité is a Human-in-the-Loop (HITL) AI application designed specifically "
    "to assist researchers in tracing African American lineages across the emancipation divide. Because pre-Civil War "
    "records are often fragmented, inconsistent, and sometimes lack surnames, identity can only be established by "
    "carefully reasoning across multiple primary sources."
)
doc.add_paragraph(
    "Verité provides a workspace where you can visually build a family tree, search normalized historical datasets "
    "(such as censuses, vital records, and Freedman's Bureau archives), and use intelligent Similarity Scoring to "
    "evaluate whether a person in a historical record is the same person in your family tree."
)

# ----------------- LAYOUT -----------------
doc.add_heading('2. Application Layout', level=1)
doc.add_paragraph(
    "The Verité interface is divided into two primary sections to facilitate simultaneous visualization and data entry:"
)
p = doc.add_paragraph(style='List Bullet')
p.add_run("Left Pane (Family Tree Canvas): ").bold = True
p.add_run("Displays the interactive visual graph of the family tree.")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Right Pane (Editors and Search): ").bold = True
p.add_run("Contains a tabbed interface for the Person Editor, Mentions Editor (Search Results), and Context (Sources).")

# ----------------- FAMILY TREE -----------------
doc.add_heading('3. Navigating the Family Tree', level=1)
doc.add_paragraph(
    "The Family Tree canvas visualizes individuals as nodes (gendered silhouettes) and relationships as connecting lines (edges). "
    "Underneath each node, the person's name and lifespan (e.g., 1850 - 1920) are displayed."
)

doc.add_heading('Interacting with the Canvas', level=2)
doc.add_paragraph("• Pan: Click and drag anywhere on the empty background to move the view.", style='List Bullet')
doc.add_paragraph("• Zoom: Use your mouse scroll wheel to zoom in and out. Double-clicking the background has been disabled to prevent accidental zooming.", style='List Bullet')
doc.add_paragraph("• Select a Person: Single-click a person node. The node will be highlighted, and the right pane will update to reflect their details.", style='List Bullet')
doc.add_paragraph("• Edit a Person: Double-click a person node to instantly open the Person Editor tab for that individual.", style='List Bullet')
doc.add_paragraph("• Context Menu: Right-click on a person node to access quick actions, such as adding a new relative or deleting the person from the tree.", style='List Bullet')

# ----------------- WORKFLOW -----------------
doc.add_heading('4. Step-by-Step Workflow', level=1)

doc.add_heading('Step A: Establishing an Anchor Person', level=2)
doc.add_paragraph(
    "Research typically begins by identifying an \"anchor\" person in a post-emancipation record, most commonly the 1880 Federal Census. "
    "Once this anchor person is created in the tree, they serve as the foundation for backward tracing."
)

doc.add_heading('Step B: Searching for Evidence', level=2)
doc.add_paragraph(
    "With a person selected in the tree, switch to the Search tab to query the historical datasets."
)
doc.add_paragraph("• You can search by name, birth year range, or specific document collections.", style='List Bullet')
doc.add_paragraph("• The system will return a list of matching Mentions (records).", style='List Bullet')

doc.add_heading('Step C: Evaluating Mentions & Similarity Scoring', level=2)
doc.add_paragraph(
    "When you click on a search result, the Mentions Editor displays a detailed breakdown of the Match Score. "
    "This score is calculated by the AI using several factors to determine confidence:"
)
doc.add_paragraph("• Name Matching: Checks for exact matches, common nicknames (e.g., \"Wm\" for William), and phonetic similarities using NYSIIS and Soundex algorithms.", style='List Bullet')
doc.add_paragraph("• Demographic Matching: Compares birth and death years, allowing for minor discrepancies common in historical documents.", style='List Bullet')
doc.add_paragraph("• Household Continuity: The most powerful factor. The system checks if the candidate record contains other family members who are already linked to your anchor person in the tree. Finding a co-resident spouse or child drastically increases the match score.", style='List Bullet')
doc.add_paragraph(
    "In the Mentions Editor detail panel, you can review the exact narrative from the record, view the individual factors contributing "
    "to the score, and see a list of Co-resident Family Members (displaying their name and birth date)."
)

doc.add_heading('Step D: Adding Evidence & Building the Tree', level=2)
doc.add_paragraph(
    "Once you are confident a record refers to your selected person, use the action buttons at the bottom of the Mentions Editor:"
)
doc.add_paragraph("• Add Mention to Person: Links the historical record to the person as verified evidence. This record will now appear in their Person Editor under 'Attached Evidence'.", style='List Bullet')
doc.add_paragraph("• See Context: Opens the original source transcription or document image for manual review.", style='List Bullet')
doc.add_paragraph("• Add Person to Tree: If the historical record lists a relative who is not yet in your family tree (e.g., you found a census record containing the anchor person's sibling), you can select the sibling in the family list and click this button to instantly create a new person node for them. Note: If the currently selected record is already attached to the active person, this button will be safely hidden to prevent duplicates.", style='List Bullet')

# ----------------- PERSON EDITOR -----------------
doc.add_heading('5. Using the Person Editor', level=1)
doc.add_paragraph(
    "The Person Editor is where you manage the master identity of an individual, synthesizing the data gathered from multiple mentions."
)
doc.add_paragraph("• Demographics: Manually adjust First, Middle, and Last names, as well as Birth/Death years. The system will auto-generate phonetic encodings to aid future searches.", style='List Bullet')
doc.add_paragraph("• Relationships: View a list of spouses, children, siblings, and parents. You can remove incorrect relationships here.", style='List Bullet')
doc.add_paragraph("• Automatic Assertion Expansion: When you add relationships, Verité handles the inverse logic automatically. If you designate Person A as the parent of Person B, the system automatically registers Person B as the child of Person A, ensuring tree consistency without double-entry.", style='List Bullet')

# ----------------- BEST PRACTICES -----------------
doc.add_heading('6. Tips and Best Practices', level=1)
doc.add_paragraph("1. Always Rely on Household Continuity: Names and ages in the 19th century were highly fluid. A strong match on co-resident family members is often more reliable than an exact spelling of a surname.", style='List Bullet')
doc.add_paragraph("2. Review Context Frequently: Use the \"See Context\" button to read the surrounding original text of a record. Often, marginalia or neighboring households provide the critical clue needed to cross the emancipation divide.", style='List Bullet')
doc.add_paragraph("3. Keep the Tree Tidy: If you accidentally add the wrong person from a mention, simply right-click their node in the tree and select \"Delete Person\".", style='List Bullet')

doc.save('c:/Bill/CC/js/AI/Verite/Verite_User_Manual.docx')
